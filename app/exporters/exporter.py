import io


def export_txt(content):
    return content.encode("utf-8")


def export_pdf(content, title):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
        leftMargin=2.5*cm, rightMargin=2.5*cm,
        topMargin=2.5*cm, bottomMargin=2.5*cm)

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("t", parent=styles["Title"],
        fontSize=18, spaceAfter=16, textColor=colors.HexColor("#1a1a2e"))
    body_style = ParagraphStyle("b", parent=styles["Normal"],
        fontSize=11, leading=17, spaceAfter=8)
    heading_style = ParagraphStyle("h", parent=styles["Heading2"],
        fontSize=13, spaceBefore=14, spaceAfter=6,
        textColor=colors.HexColor("#0066cc"))

    story = [Paragraph(title, title_style), Spacer(1, 0.3*cm)]

    for line in content.split("\n"):
        line = line.strip()
        if not line:
            story.append(Spacer(1, 0.2*cm))
        elif line.startswith("## "):
            safe_h = line[3:].replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
            story.append(Paragraph(safe_h, heading_style))
        elif line.startswith("# "):
            safe_h = line[2:].replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
            story.append(Paragraph(safe_h, title_style))
        elif line.startswith("- ") or line.startswith("* "):
            safe_b = line[2:].replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
            story.append(Paragraph(f"• {safe_b}", body_style))
        else:
            safe = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            story.append(Paragraph(safe, body_style))

    doc.build(story)
    return buf.getvalue()


def export_docx(content, title):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    heading = doc.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for line in content.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph("")
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
